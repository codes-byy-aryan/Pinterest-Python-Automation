"""
Pinterest → Gemini → WhatsApp Daily Automation
Runs every day at 10:00 AM via cron.
"""

import os
import sys
import logging
import json
import time
import requests
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

# ── Load .env ─────────────────────────────────────────────────────────────────
load_dotenv()

# ── Paths ─────────────────────────────────────────────────────────────────────
BASE_DIR      = Path(__file__).parent
KEYWORDS_FILE = BASE_DIR / "keywords.txt"
IDEAS_FILE    = BASE_DIR / "ideas.docx"
LOG_FILE      = BASE_DIR / "automation.log"

# ── Logging ───────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler(LOG_FILE),
        logging.StreamHandler(sys.stdout),
    ],
)
log = logging.getLogger(__name__)

# ── Env vars ──────────────────────────────────────────────────────────────────
GEMINI_API_KEY       = os.getenv("GEMINI_API_KEY")
WHATSAPP_TOKEN       = os.getenv("WHATSAPP_TOKEN")
WHATSAPP_PHONE_ID    = os.getenv("WHATSAPP_PHONE_ID")
WHATSAPP_TO_NUMBER   = os.getenv("WHATSAPP_TO_NUMBER")   # e.g. 919876543210
PINTEREST_ACCESS_TOKEN = os.getenv("PINTEREST_ACCESS_TOKEN")  # optional

# ─────────────────────────────────────────────────────────────────────────────
# STEP 1 — Collect Pinterest Trending Keywords
# ─────────────────────────────────────────────────────────────────────────────

# Broad evergreen seed topics that generate fresh keyword combos each day.
# If you have a Pinterest API token, the script will call the real API first.
SEED_TOPICS = [
    "home decor", "wedding ideas", "fitness motivation", "healthy recipes",
    "hair styles", "nail art", "fashion outfits", "travel destinations",
    "DIY crafts", "minimalist design", "skincare routine", "budget living",
    "boho aesthetic", "vision board", "mental health quotes",
    "small business ideas", "digital art", "cute animals", "morning routines",
    "meal prep", "plant decor", "journal ideas", "photography tips",
    "book recommendations", "tattoo ideas",
]

SEASONAL_KEYWORDS = [
    f"{datetime.now().strftime('%B')} fashion",
    f"{datetime.now().strftime('%B')} recipes",
    f"{datetime.now().year} trends",
    "summer outfits", "fall home decor", "winter skincare", "spring garden",
    "holiday gift ideas", "new year goals", "back to school",
]

NICHE_KEYWORDS = [
    "aesthetic room decor", "cottagecore outfits", "dark academia style",
    "Y2K fashion", "clean girl aesthetic", "mob wife aesthetic",
    "quiet luxury fashion", "coastal grandmother style", "indie room ideas",
    "Pinterest board ideas", "vision board 2025", "manifestation journal",
    "self care Sunday", "glow up tips", "blonde highlights ideas",
    "french manicure designs", "coquette aesthetic", "balletcore outfit",
    "old money style", "Pinterest worthy",
]


def fetch_pinterest_keywords_api() -> list[str]:
    """Try to get trends from Pinterest API (requires access token)."""
    if not PINTEREST_ACCESS_TOKEN:
        return []
    try:
        headers = {"Authorization": f"Bearer {PINTEREST_ACCESS_TOKEN}"}
        url = "https://api.pinterest.com/v5/trends/keywords"
        params = {"region": "US", "trend_type": "growing", "limit": 50}
        r = requests.get(url, headers=headers, params=params, timeout=15)
        if r.status_code == 200:
            data = r.json()
            keywords = [item.get("keyword", "") for item in data.get("trends", []) if item.get("keyword")]
            if keywords:
                log.info(f"✅ Fetched {len(keywords)} keywords from Pinterest API.")
                return keywords
    except Exception as e:
        log.warning(f"Pinterest API call failed: {e}")
    return []


def collect_keywords() -> list[str]:
    """Collect 50–100 Pinterest-style trending keywords."""
    log.info("🔍 Collecting Pinterest keywords...")
    api_keywords = fetch_pinterest_keywords_api()
    combined = list(dict.fromkeys(
        api_keywords + SEED_TOPICS + SEASONAL_KEYWORDS + NICHE_KEYWORDS
    ))
    keywords = combined[:100]
    KEYWORDS_FILE.write_text("\n".join(keywords), encoding="utf-8")
    log.info(f"✅ Saved {len(keywords)} keywords → {KEYWORDS_FILE}")
    return keywords


# ─────────────────────────────────────────────────────────────────────────────
# STEP 2 — Send to Gemini & Generate Ideas
# ─────────────────────────────────────────────────────────────────────────────

GEMINI_URL = (
    "https://generativelanguage.googleapis.com/v1beta/models/"
    "gemini-1.5-flash:generateContent"
)

GEMINI_PROMPT_TEMPLATE = """
You are a Pinterest content strategist. For each keyword below, generate a structured content brief.
Format your response as a JSON array. Each object must have exactly these fields:
- keyword
- content_ideas   (list of 3 ideas)
- titles          (list of 3 pin titles)
- topics          (list of 3 topics)
- description     (2-sentence pin description)
- requirements    (list of what images/assets are needed)
- hashtags        (list of 8 relevant hashtags)

Return ONLY valid JSON, no markdown fences, no extra text.

Keywords:
{keywords}
"""


def call_gemini(keywords: list[str]) -> list[dict]:
    """Call Gemini API and return parsed ideas list."""
    if not GEMINI_API_KEY:
        raise ValueError("GEMINI_API_KEY is not set in .env")

    log.info("🤖 Sending keywords to Gemini API...")

    # Process in batches of 10 to stay within token limits
    all_ideas = []
    batch_size = 10

    for i in range(0, len(keywords), batch_size):
        batch = keywords[i:i + batch_size]
        prompt = GEMINI_PROMPT_TEMPLATE.format(keywords="\n".join(batch))

        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.8,
                "maxOutputTokens": 8192,
            },
        }
        params = {"key": GEMINI_API_KEY}

        resp = requests.post(GEMINI_URL, json=payload, params=params, timeout=60)
        resp.raise_for_status()

        raw = resp.json()
        text = raw["candidates"][0]["content"]["parts"][0]["text"]

        # Strip markdown fences if present
        text = text.strip()
        if text.startswith("```"):
            text = "\n".join(text.split("\n")[1:])
        if text.endswith("```"):
            text = "\n".join(text.split("\n")[:-1])

        try:
            batch_ideas = json.loads(text)
            all_ideas.extend(batch_ideas)
            log.info(f"  ✅ Batch {i//batch_size + 1}: {len(batch_ideas)} ideas generated.")
        except json.JSONDecodeError as e:
            log.warning(f"  ⚠️ JSON parse error in batch {i//batch_size + 1}: {e}")
            # Fallback: store raw text
            for kw in batch:
                all_ideas.append({"keyword": kw, "raw_response": text})

        time.sleep(1)  # Rate-limit courtesy

    log.info(f"✅ Gemini returned ideas for {len(all_ideas)} keywords.")
    return all_ideas


# ─────────────────────────────────────────────────────────────────────────────
# STEP 3 — Save ideas.docx
# ─────────────────────────────────────────────────────────────────────────────

def build_docx(ideas: list[dict]) -> None:
    """Create a formatted ideas.docx from Gemini output."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    log.info("📄 Building ideas.docx...")

    doc = Document()

    # Title
    title = doc.add_heading("Pinterest Content Ideas", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0xE6, 0x00, 0x23)  # Pinterest red

    date_para = doc.add_paragraph(f"Generated: {datetime.now().strftime('%d %B %Y, %I:%M %p')}")
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()  # spacer

    for idx, item in enumerate(ideas, 1):
        kw = item.get("keyword", "Unknown")

        # Section heading
        h = doc.add_heading(f"{idx}. {kw.title()}", level=1)
        h.runs[0].font.color.rgb = RGBColor(0xE6, 0x00, 0x23)

        if "raw_response" in item:
            doc.add_paragraph(item["raw_response"])
            doc.add_paragraph()
            continue

        def add_field(label: str, value):
            p = doc.add_paragraph()
            run_label = p.add_run(f"{label}: ")
            run_label.bold = True
            run_label.font.size = Pt(11)
            if isinstance(value, list):
                run_val = p.add_run(", ".join(str(v) for v in value))
            else:
                run_val = p.add_run(str(value))
            run_val.font.size = Pt(11)

        add_field("💡 Content Ideas",  item.get("content_ideas", []))
        add_field("📌 Pin Titles",     item.get("titles", []))
        add_field("🗂️  Topics",         item.get("topics", []))
        add_field("📝 Description",    item.get("description", ""))
        add_field("🛠️  Requirements",   item.get("requirements", []))
        add_field("🏷️  Hashtags",       item.get("hashtags", []))

        doc.add_paragraph()

    doc.save(str(IDEAS_FILE))
    log.info(f"✅ ideas.docx saved → {IDEAS_FILE}")


# ─────────────────────────────────────────────────────────────────────────────
# STEP 4 — Send via WhatsApp Cloud API
# ─────────────────────────────────────────────────────────────────────────────

def upload_media_to_whatsapp() -> str:
    """Upload ideas.docx to WhatsApp media endpoint and return media_id."""
    url = f"https://graph.facebook.com/v19.0/{WHATSAPP_PHONE_ID}/media"
    headers = {"Authorization": f"Bearer {WHATSAPP_TOKEN}"}
    with open(IDEAS_FILE, "rb") as f:
        files = {
            "file": ("ideas.docx", f,
                     "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            "messaging_product": (None, "whatsapp"),
        }
        r = requests.post(url, headers=headers, files=files, timeout=30)
    r.raise_for_status()
    media_id = r.json()["id"]
    log.info(f"✅ Uploaded ideas.docx to WhatsApp. Media ID: {media_id}")
    return media_id


def send_whatsapp(media_id: str) -> None:
    """Send the uploaded document to the WhatsApp number."""
    url = f"https://graph.facebook.com/v19.0/{WHATSAPP_PHONE_ID}/messages"
    headers = {
        "Authorization": f"Bearer {WHATSAPP_TOKEN}",
        "Content-Type": "application/json",
    }
    today = datetime.now().strftime("%d %B %Y")
    payload = {
        "messaging_product": "whatsapp",
        "to": WHATSAPP_TO_NUMBER,
        "type": "document",
        "document": {
            "id": media_id,
            "filename": "ideas.docx",
            "caption": (
                f"📌 *Pinterest Content Ideas — {today}*\n\n"
                "Your daily AI-generated Pinterest content brief is ready! "
                "Keywords, titles, descriptions, hashtags & more. 🚀"
            ),
        },
    }
    r = requests.post(url, headers=headers, json=payload, timeout=30)
    r.raise_for_status()
    log.info(f"✅ ideas.docx sent to WhatsApp number: {WHATSAPP_TO_NUMBER}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

def main():
    log.info("=" * 60)
    log.info(f"🚀 Pinterest Automation started — {datetime.now()}")
    log.info("=" * 60)

    try:
        # 1. Keywords
        keywords = collect_keywords()

        # 2. Gemini
        ideas = call_gemini(keywords)

        # 3. DOCX
        build_docx(ideas)

        # 4. WhatsApp
        if not all([WHATSAPP_TOKEN, WHATSAPP_PHONE_ID, WHATSAPP_TO_NUMBER]):
            log.warning("⚠️  WhatsApp env vars missing — skipping send.")
        else:
            media_id = upload_media_to_whatsapp()
            send_whatsapp(media_id)

        log.info("🎉 Automation completed successfully!")

    except Exception as e:
        log.error(f"❌ Automation failed: {e}", exc_info=True)
        sys.exit(1)


if __name__ == "__main__":
    main()
